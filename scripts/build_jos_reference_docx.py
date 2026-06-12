#!/usr/bin/env python3
"""
Build a Pandoc reference DOCX template matching 《软件学报》2025 format.

Generates scripts/pandoc_filters/jos_reference.docx with correct page size,
margins, fonts, and heading styles extracted from the journal template specs.
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(rpr, ascii_name, east_asia_name, size_pt, bold=False):
    """Set font properties on a run-properties element."""
    fonts = rpr.find(qn('w:rFonts'))
    if fonts is None:
        fonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, fonts)
    fonts.set(qn('w:ascii'), ascii_name)
    fonts.set(qn('w:hAnsi'), ascii_name)
    fonts.set(qn('w:eastAsia'), east_asia_name)

    sz = rpr.find(qn('w:sz'))
    if sz is None:
        sz = rpr.makeelement(qn('w:sz'), {})
        rpr.append(sz)
    sz.set(qn('w:val'), str(int(size_pt * 2)))  # half-points

    szCs = rpr.find(qn('w:szCs'))
    if szCs is None:
        szCs = rpr.makeelement(qn('w:szCs'), {})
        rpr.append(szCs)
    szCs.set(qn('w:val'), str(int(size_pt * 2)))

    if bold:
        b = rpr.find(qn('w:b'))
        if b is None:
            b = rpr.makeelement(qn('w:b'), {})
            rpr.append(b)
    else:
        b = rpr.find(qn('w:b'))
        if b is not None:
            rpr.remove(b)

def set_spacing(ppr, before_pt=0, after_pt=0, line_twips=None):
    """Set paragraph spacing."""
    spacing = ppr.find(qn('w:spacing'))
    if spacing is None:
        spacing = ppr.makeelement(qn('w:spacing'), {})
        ppr.append(spacing)
    spacing.set(qn('w:before'), str(int(before_pt * 20)))  # twips
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    if line_twips is not None:
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')

def set_indent(ppr, first_line_pt=None):
    """Set paragraph indentation."""
    ind = ppr.find(qn('w:ind'))
    if ind is None:
        ind = ppr.makeelement(qn('w:ind'), {})
        ppr.append(ind)
    if first_line_pt is not None:
        ind.set(qn('w:firstLine'), str(int(first_line_pt * 20)))

def main():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(18.40)
    section.page_height = Cm(26.00)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.top_margin = Cm(1.00)
    section.bottom_margin = Cm(2.20)
    section.header_distance = Cm(0.50)
    section.footer_distance = Cm(1.00)

    # ── Modify styles ──
    styles = doc.styles

    # -- Normal (正文): 宋体 9pt, single spacing, 2em indent --
    normal = styles['Normal']
    normal_ppr = normal.element.find(qn('w:pPr'))
    if normal_ppr is None:
        normal_ppr = normal.element.makeelement(qn('w:pPr'), {})
        normal.element.append(normal_ppr)

    normal_rpr = normal.element.find(qn('w:rPr'))
    if normal_rpr is None:
        normal_rpr = normal.element.makeelement(qn('w:rPr'), {})
        normal.element.append(normal_rpr)

    set_font(normal_rpr, 'Times New Roman', '宋体', 9)
    set_spacing(normal_ppr, before_pt=0, after_pt=0, line_twips=260)  # ~13pt
    set_indent(normal_ppr, first_line_pt=18)  # 2em at 9pt ≈ 18pt

    # -- Title: 黑体 14pt bold --
    title_style = styles['Title']
    title_ppr = title_style.element.find(qn('w:pPr'))
    if title_ppr is None:
        title_ppr = title_style.element.makeelement(qn('w:pPr'), {})
        title_style.element.append(title_ppr)
    title_rpr = title_style.element.find(qn('w:rPr'))
    if title_rpr is None:
        title_rpr = title_style.element.makeelement(qn('w:rPr'), {})
        title_style.element.append(title_rpr)
    set_font(title_rpr, 'Times New Roman', '黑体', 14, bold=True)
    set_spacing(title_ppr, before_pt=6, after_pt=6)
    # Remove first line indent for title
    ind = title_ppr.find(qn('w:ind'))
    if ind is not None:
        ind.attrib.pop(qn('w:firstLine'), None)

    # -- Heading 1: 黑体 10.5pt (五号), space before/after 8pt --
    h1 = styles['Heading 1']
    h1_ppr = h1.element.find(qn('w:pPr'))
    if h1_ppr is None:
        h1_ppr = h1.element.makeelement(qn('w:pPr'), {})
        h1.element.append(h1_ppr)
    h1_rpr = h1.element.find(qn('w:rPr'))
    if h1_rpr is None:
        h1_rpr = h1.element.makeelement(qn('w:rPr'), {})
        h1.element.append(h1_rpr)
    set_font(h1_rpr, 'Times New Roman', '黑体', 10.5, bold=True)
    set_spacing(h1_ppr, before_pt=8, after_pt=8)
    # Remove indent for headings
    ind = h1_ppr.find(qn('w:ind'))
    if ind is not None:
        ind.attrib.pop(qn('w:firstLine'), None)

    # -- Heading 2: 黑体 9pt (小五号), space before/after 1.2pt --
    h2 = styles['Heading 2']
    h2_ppr = h2.element.find(qn('w:pPr'))
    if h2_ppr is None:
        h2_ppr = h2.element.makeelement(qn('w:pPr'), {})
        h2.element.append(h2_ppr)
    h2_rpr = h2.element.find(qn('w:rPr'))
    if h2_rpr is None:
        h2_rpr = h2.element.makeelement(qn('w:rPr'), {})
        h2.element.append(h2_rpr)
    set_font(h2_rpr, 'Times New Roman', '黑体', 9, bold=True)
    set_spacing(h2_ppr, before_pt=1.2, after_pt=1.2)
    ind = h2_ppr.find(qn('w:ind'))
    if ind is not None:
        ind.attrib.pop(qn('w:firstLine'), None)

    # -- Heading 3 --
    h3 = styles['Heading 3']
    h3_ppr = h3.element.find(qn('w:pPr'))
    if h3_ppr is None:
        h3_ppr = h3.element.makeelement(qn('w:pPr'), {})
        h3.element.append(h3_ppr)
    h3_rpr = h3.element.find(qn('w:rPr'))
    if h3_rpr is None:
        h3_rpr = h3.element.makeelement(qn('w:rPr'), {})
        h3.element.append(h3_rpr)
    set_font(h3_rpr, 'Times New Roman', '黑体', 9, bold=True)
    set_spacing(h3_ppr, before_pt=1, after_pt=1)

    # -- Abstract / Block Text style --
    # Pandoc maps abstract to "Abstract" style or uses Block Text
    for sname in ['Abstract', 'Block Text']:
        try:
            s = styles[sname]
        except KeyError:
            continue
        s_ppr = s.element.find(qn('w:pPr'))
        if s_ppr is None:
            s_ppr = s.element.makeelement(qn('w:pPr'), {})
            s.element.append(s_ppr)
        s_rpr = s.element.find(qn('w:rPr'))
        if s_rpr is None:
            s_rpr = s.element.makeelement(qn('w:rPr'), {})
            s.element.append(s_rpr)
        set_font(s_rpr, 'Times New Roman', '楷体', 9)
        set_spacing(s_ppr, before_pt=0, after_pt=0, line_twips=240)

    # -- Figure caption --
    try:
        cap = styles['Caption']
        cap_ppr = cap.element.find(qn('w:pPr'))
        if cap_ppr is None:
            cap_ppr = cap.element.makeelement(qn('w:pPr'), {})
            cap.element.append(cap_ppr)
        cap_rpr = cap.element.find(qn('w:rPr'))
        if cap_rpr is None:
            cap_rpr = cap.element.makeelement(qn('w:rPr'), {})
            cap.element.append(cap_rpr)
        set_font(cap_rpr, 'Times New Roman', '宋体', 9)
        # center alignment
        jc = cap_ppr.find(qn('w:jc'))
        if jc is None:
            jc = cap_ppr.makeelement(qn('w:jc'), {})
            cap_ppr.append(jc)
        jc.set(qn('w:val'), 'center')
    except KeyError:
        pass

    # Add minimal content so pandoc can detect the styles
    doc.add_paragraph('Title placeholder', style='Title')
    doc.add_heading('Heading 1', level=1)
    doc.add_heading('Heading 2', level=2)
    doc.add_paragraph('Body text paragraph.', style='Normal')

    outpath = 'scripts/pandoc_filters/jos_reference.docx'
    doc.save(outpath)
    print(f'Reference template saved to {outpath}')

if __name__ == '__main__':
    main()
