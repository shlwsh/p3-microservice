#!/usr/bin/env python3
"""Extract formatting details from 2025 JOS template doc."""
import docx
from docx.shared import Pt, Cm

doc = docx.Document('docs/latex-models/software-journal/软件学报排版样例2025年版.doc')

# Section/page setup
for i, section in enumerate(doc.sections):
    print(f"=== Section {i} ===")
    print(f"  Page: {section.page_width.cm:.2f} x {section.page_height.cm:.2f} cm")
    print(f"  Margins L/R/T/B: {section.left_margin.cm:.2f} / {section.right_margin.cm:.2f} / {section.top_margin.cm:.2f} / {section.bottom_margin.cm:.2f} cm")
    print(f"  Header/Footer dist: {section.header_distance.cm:.2f} / {section.footer_distance.cm:.2f} cm")
    print(f"  Gutter: {section.gutter.cm:.2f} cm")

# Headers
for i, section in enumerate(doc.sections):
    header = section.header
    if header:
        print(f"\n=== Header {i} ===")
        for p in header.paragraphs:
            for run in p.runs:
                sz = f"{run.font.size.pt:.1f}" if run.font.size else '-'
                print(f'  [{run.font.name}/{sz}pt/b={run.font.bold}/i={run.font.italic}] "{run.text[:60]}"')

# Paragraphs with format
print("\n=== Paragraphs ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    pf = p.paragraph_format
    style = p.style.name if p.style else '-'
    align = str(pf.alignment) if pf.alignment else '-'
    first_indent = f"{pf.first_line_indent.pt:.1f}pt" if pf.first_line_indent else '-'
    space_before = f"{pf.space_before.pt:.1f}pt" if pf.space_before else '-'
    space_after = f"{pf.space_after.pt:.1f}pt" if pf.space_after else '-'
    line_sp = str(pf.line_spacing) if pf.line_spacing else '-'

    print(f'\nP{i} [{style}] align={align} first_indent={first_indent} sp_b={space_before} sp_a={space_after} line_sp={line_sp}')
    print(f'  >> {text[:100]}')

    for run in p.runs:
        if run.text.strip():
            f = run.font
            sz = f"{f.size.pt:.1f}" if f.size else '-'
            print(f'  RUN: {f.name}/{sz}pt b={f.bold} i={f.italic} "{run.text[:50]}"')

# Tables
print(f"\n=== Tables: {len(doc.tables)} ===")
for i, table in enumerate(doc.tables):
    print(f"Table {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    for j, row in enumerate(table.rows):
        for k, cell in enumerate(row.cells):
            print(f"  [{j},{k}] {cell.text[:40]}")
